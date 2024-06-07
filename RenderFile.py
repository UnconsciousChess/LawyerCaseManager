# 导入自带模块-时间
import time
# 导入自带模块-系统
import os,sys

# 导入第三方库docxtpl
from docxtpl import DocxTemplate

# 导入第三方库python-docx
from docx import Document
from docx.shared import Pt    # 字体大小转换模块


def DeleteFileIfExist(OutputDir,FileName) -> None:
    # 判断是否存在同名文件，如果存在就删除原文件
    if os.path.exists(OutputDir + "\\" + FileName):
        # 防止因文件被WINWORD或其他应用打开中而删除失败
        while True:
            try:
                os.remove(OutputDir + "\\" + FileName)
                break
            except:
                time.sleep(3)
                print("该文件被其他程序占用，删除原同名文件失败,等待3秒后重试")
    return
    

def RenderFileInDocxtpl(TemplateFileDir,Case,OutputDir) -> None:
        
    # 读取模板文件
    doc = DocxTemplate(TemplateFileDir)
    # 获取模板文件名
    TemplaterFileName = TemplateFileDir.split("\\")[-1]
    
    # 初始化文件是否在我方当事人需要分开的列表和对方当事人需要分开的列表
    CurrentTemplateFileIsInOurClientMultipleFileList = False
    CurrentTemplateFileIsInOppositeMultipleFileList = False

    context = {
        # 自动生成时间信息
        '年' : time.strftime("%Y",time.localtime()),
        '月' : time.strftime("%m",time.localtime()),
        '日' : time.strftime("%d",time.localtime()),
    }

    # 案件共同信息（非当事人信息）
    context["案由"] = Case.GetCauseOfAction()
    context["管辖法院"] = Case.GetJurisdiction()
    context["委托阶段"] = Case.GetCaseAgentStageStr()

    # 如果委托阶段包括立案阶段，那么案号不可能存在，应当写作“本案未立案，最终以实际案号为准”
    if 1 in Case.GetCaseAgentStage():
        context["案号"] = "本案未立案，最终以实际案号为准"
    else:
        context["案号"] = Case.GetCaseCourtCode()

    # 获取全部原告信息
    context["全部原告名称"] = Case.GetAllPlaintiffNames()
    # 获取全部被告信息
    context["全部被告名称"] = Case.GetAllDefendantNames()
    # 委托付费信息

    # 先判断是否需要收费,如果status是None，则不需要收费；是True则为风险收费；是False则为固定收费
    if Case.GetRiskAgentStatus() is not None:
        # 先判断是否为风险收费
        if Case.GetRiskAgentStatus() == True:                 # 风险收费
            context["风险代理前期律师费"] = Case.GetRiskAgentUpfrontFee()
            context["风险代理后期比例"] = Case.GetRiskAgentPostFeeRate()
        if Case.GetRiskAgentStatus() == False:                # 固定收费
            FixedFeeRuleStr = ""
            i = 0
            # 先判断案件阶段与案件收费的长度是否一致
            if len(Case.GetAgentFixedFee()) == len(Case.GetCaseAgentStage()):
                if Case.GetAgentFixedFee() != []:
                    for stage in Case.GetCaseAgentStage():
                        if stage == 1:
                            FixedFeeRuleStr += "一审立案阶段收费：" + str(Case.GetAgentFixedFee()[i]) + "元；"
                            i += 1
                        elif stage == 2:
                            FixedFeeRuleStr += "一审审理阶段收费：" + str(Case.GetAgentFixedFee()[i]) + "元；"
                            i += 1
                        elif stage == 3:
                            FixedFeeRuleStr += "二审阶段收费：" + str(Case.GetAgentFixedFee()[i]) + "元；"
                            i += 1                  
                        elif stage == 4:
                            FixedFeeRuleStr += "执行阶段收费" + str(Case.GetAgentFixedFee()[i]) + "元；"
                            i += 1
                        elif stage == 5:
                            FixedFeeRuleStr += "再审阶段：" + str(Case.GetAgentFixedFee()[i]) + "元；"
                            i += 1    
            # 最后赋值到RenderMaterial字典中的FixedAgentFeeRuleStr键里面
            context["固定代理费收费规则"] = FixedFeeRuleStr


    # 获取我方当事人列表
    OurClientList,OurClientSide = Case.GetOurClientListAndSide()
    # 如果代理原告,对方全部当事人则为被告
    if OurClientSide == "p":
        context['对方全部当事人'] = context["全部被告名称"]
        OppositeSideLitigantList = Case.GetDefendantList()
    # 如果代理被告，对方全部当事人则为原告
    if OurClientSide == "d":
        context['对方全部当事人'] = context["全部原告名称"]
        OppositeSideLitigantList = Case.GetPlaintiffList()
    
    # 得到我方当事人的字符串AllClients，以顿号分割
    AllClients = "" 
    for client in OurClientList:
        AllClients += client.GetName() + "、"
    context['我方全部当事人'] = AllClients[:-1]    # 去掉最后一个顿号

    # 判断是否为需要分开不同人做的材料
           
    # 我方当事人需要分开的列表
    OurClientMultipleFileList = ["授权委托书","征求意见表","账户确认书","提交材料清单","地址确认书"]     
    # 对方当事人需要分开的列表
    OppositeMultipleFileList = ["线索书"]

    # 先判断本模板是否在【我方当事人需要分开的列表】
    for name in OurClientMultipleFileList:
        if name in TemplaterFileName:
            CurrentTemplateFileIsInOurClientMultipleFileList = True
    # 再判断本模板是否在【对方当事人需要分开的列表】
    for name in OppositeMultipleFileList:
        if name in TemplaterFileName:
            CurrentTemplateFileIsInOppositeMultipleFileList = True

    # 如果当前文书是我方当事人需要分开的文书，则遍历我方当事人列表并分别生成
    if CurrentTemplateFileIsInOurClientMultipleFileList:
        for client in OurClientList:
            # 如果我方当事人是法人或者其他组织
            if client.GetLitigantType() == 2 or client.GetLitigantType() == 3:
                # 如果该文书是自然人的模版，则跳过该当事人
                if "自然人版"  in TemplaterFileName:
                    continue
                # 下面获取法人代表名称
                context["我方当事人法定代表人"] = client.GetLegalRepresentative()
                # 获取法人代表身份证号码
                context["我方当事人法定代表人身份号码"] = client.GetLegalRepresentativeIDCode()
            #  如果我方当事人是自然人
            if client.GetLitigantType() == 1:
                if "法人、其他组织版"  in TemplaterFileName:
                    continue
            # 填入context
            context['我方当事人名称'] = client.GetName()
            context['我方当事人身份号码'] = client.GetIDCode()
            context['我方当事人地址'] = client.GetLocation()
            context['我方当事人电话'] = client.GetContactNumber()
            context['我方当事人银行账户名'] = client.GetBankAccount().GetAccountName()
            context['我方当事人开户行'] = client.GetBankAccount().GetBankName()
            context['我方当事人银行账户号码'] = client.GetBankAccount().GetAccountNumber()

            # 渲染模板
            doc.render(context)
            DocSaveName = TemplaterFileName.replace(".docx", "（" + client.GetName() + "）.docx")
            # 判断是否存在同名文件，如果存在就删除原文件
            DeleteFileIfExist(OutputDir,DocSaveName)
            # 保存文件
            doc.save(OutputDir + "\\" + DocSaveName)

    # 如果当前文书是对方当事人需要分开的文书，则遍历对方当事人列表并分别生成
    if CurrentTemplateFileIsInOppositeMultipleFileList:
        for opposlitigant in OppositeSideLitigantList:
            # 如果该对方当事人是法人或者其他组织，则需要
            if opposlitigant.GetLitigantType() == 2 or opposlitigant.GetLitigantType() == 3:
                # 如果该文书是自然人的模版，则跳过该当事人
                if "自然人版"  in TemplaterFileName:
                    continue
                    # 下面获取法人代表名称
                context["对方当事人法定代表人"] = opposlitigant.GetLegalRepresentative()
                # 获取法人代表身份证号码
                context["对方当事人法定代表人身份号码"] = opposlitigant.GetLegalRepresentativeIDCode()                   
                #  如果我方当事人是自然人
            if opposlitigant.GetLitigantType() == 1:
                if "法人、其他组织版"  in TemplaterFileName:
                    continue
            # 填入context
            context['对方当事人名称'] = opposlitigant.GetName()
            context['对方当事人身份号码'] = opposlitigant.GetIDCode()
            context['对方当事人地址'] = opposlitigant.GetLocation()
            context['对方当事人电话'] = opposlitigant.GetContactNumber()

            # 渲染模板
            doc.render(context)
            DocSaveName = TemplaterFileName.replace(".docx", "（" + opposlitigant.GetName() + "）.docx")
            # 判断是否存在同名文件，如果存在就删除原文件
            DeleteFileIfExist(OutputDir,DocSaveName)
            # 保存文件
            doc.save(OutputDir + "\\" + DocSaveName)

    # 如果当前文书不需要分开，则直接渲染
    if (not CurrentTemplateFileIsInOurClientMultipleFileList and 
        not CurrentTemplateFileIsInOppositeMultipleFileList):
        # 直接渲染模板
        doc.render(context)
        # 判断是否存在同名文件，如果存在就删除原文件
        DeleteFileIfExist(OutputDir,TemplaterFileName)
        # 保存文件
        doc.save(OutputDir + "\\" + TemplaterFileName)

    return
              

def RenderFileInDOCX(TemplateFileDir,Case,OutputDir) -> None:
    TemplateFileName = TemplateFileDir.split("\\")[-1]
    # 如果TemplateFileDir的文件名为起诉状，则用下面的代码
    if "起诉状" in TemplateFileName:
        # 实例化一个Document对象
        Doc = Document(TemplateFileDir)
        for para in Doc.paragraphs:
            # 填写诉讼参与人信息
            if "LitigantInformation" in para.text:
                # 先删除该行
                para.text = ""
                para.style.font.size = Pt(12)
                # 写入所有原告信息
                for plaintiff in Case.GetPlaintiffList():
                    # 如果只有一个原告，就不用写序号
                    if len(Case.GetPlaintiffList()) == 1:
                        plaintiffinfo = "原告：" + plaintiff.GetName() + "，"
                    else:
                        plaintiffinfo = "原告" + str(Case.GetPlaintiffList().index(plaintiff)+1) + ":" + plaintiff.GetName() + "，"
                    if plaintiff.GetLitigantType() == 1:
                        plaintiffinfo +=  "身份证号码：" + plaintiff.GetIDCode() + "，" 
                    elif plaintiff.GetLitigantType() == 2 or plaintiff.GetLitigantType() == 3:
                        plaintiffinfo += "统一社会信用代码：" + plaintiff.GetIDCode() + "，"
                        plaintiffinfo += "法定代表人/负责人：" + plaintiff.GetLegalRepresentative() + "，"
                    # 地址和联系方式
                    if plaintiff.GetLocation() != None:
                        plaintiffinfo += "地址：" + plaintiff.GetLocation() + "，"
                    if plaintiff.GetContactNumber() != None:
                        plaintiffinfo += "联系方式：" + plaintiff.GetContactNumber() + "，"
                    # 删掉最后一个逗号
                    plaintiffinfo = plaintiffinfo[:-1]

                    run = para.add_run(plaintiffinfo + "\n")
                
                #原被告之间空一行
                para.add_run("\n")

                # 写入所有被告信息
                for defendant in Case.GetDefendantList():
                    # 如果只有一个被告，就不用写序号
                    if len(Case.GetDefendantList()) == 1:
                        defendantinfo = "被告：" + defendant.GetName() + "，"
                    else:    
                        defendantinfo = "被告" + str(Case.GetDefendantList().index(defendant)+1) + ":" + defendant.GetName() + "，"
                    if defendant.GetLitigantType() == 1:
                        defendantinfo +=  "身份证号码：" + defendant.GetIDCode() + "，" 
                    elif defendant.GetLitigantType() == 2 or defendant.GetLitigantType() == 3:
                        defendantinfo += "统一社会信用代码：" + defendant.GetIDCode() + "，"
                        defendantinfo += "法定代表人/负责人：" + defendant.GetLegalRepresentative() + "，"
                    # 地址和联系方式
                    if defendant.GetLocation() != None:
                        defendantinfo += "地址：" + defendant.GetLocation() + "，"
                    if defendant.GetContactNumber() != None:
                        defendantinfo += "联系方式：" + defendant.GetContactNumber() + "，"
                    # 删掉最后一个逗号
                    defendantinfo = defendantinfo[:-1]

                    run = para.add_run(defendantinfo + "\n")

            
            # 替换案由
            if "CauseOfAction" in para.text:
                para.text = para.text.replace("CauseOfAction",Case.GetCauseOfAction())
                para.style.font.size = Pt(12)
            # 替换此致后面的管辖法院
            if "CourtName" in para.text:
                para.text = para.text.replace("CourtName",Case.GetJurisdiction())
                para.style.font.size = Pt(12)
            # 写每个原告的姓名，每个空一行
            if "PlaintiffSignature" in para.text:
                # 把PlaintiffSignature删掉
                para.text = para.text.replace("PlaintiffSignature","")
                # 当事人为自然人时，在后面加上（签名）
                for plaintiff in Case.GetPlaintiffList():
                    count = 0
                    if plaintiff.GetLitigantType() == 1:
                        para.add_run(plaintiff.GetName() + "（签名）" )
                        count += 1
                    if plaintiff.GetLitigantType() == 2 or plaintiff.GetLitigantType() == 3:
                        para.add_run(plaintiff.GetName() + "（盖章）" )
                        count += 1
                    # 如果不是最后一个原告，则加一个换行
                    if count != len(Case.GetPlaintiffList()):
                        para.add_run("\n")
            # 落款后写日期
            if "SignatureDate" in para.text:
                para.text = para.text.replace("SignatureDate",time.strftime("%Y年%m月%d日",time.localtime())) 
                para.style.font.size = Pt(12)

        # 去掉文件名后缀
        FileName = TemplateFileName.replace(".docx","")
        # 加上原被告双方的名字，中间用"V."隔开，再加上后缀
        FileName += "（" + Case.GetAllPlaintiffNames() + "v." + Case.GetAllDefendantNames() + "）" + ".docx"
        # 检查同名文件并删除
        DeleteFileIfExist(OutputDir,FileName)           
        # 保存文件
        Doc.save(FileName)


# 归档卷内目录自动生成功能
def RenderArchiveDirectory(TemplateFileDir,RenderDict,OutputDir) -> None:
    # 获取模板文件名
    TemplateFileName = TemplateFileDir.split("\\")[-1]
    # 实例化模板文件
    Doc = Document(TemplateFileDir)
    
    # 获取其中的表格,因为本文件只有一个表格，所以直接取第一个表格
    table = Doc.tables[0]
    
    # 目前没有发现这个包里面有直接获取对应单元格的行数和列数的方法，所以只能用下面的方法
    # 遍历表格中第一列的所有单元格，并将其行数储存在一个字典中
    RowNumberDict = {}
    RowNumber = 0
    for cell in table.columns[0].cells:
        # 单元格的文字作为字典的key，单元格的行数作为字典的value
        RowNumberDict[cell.text] = RowNumber
        RowNumber += 1

    # 对比两个字典，如果RenderDict的key在RowNumberDict的key中，说明该行属性是要渲染的，则按下面的规则进行
    for Tablekey in RowNumberDict.keys():
        for RenderDictkey in RenderDict.keys():
            # 如果表格中的某个单元格的文字与字典中的key相同
            if RenderDictkey in Tablekey:
                # 如果字典中的key对应的value的第一个元素为True，为原件
                if RenderDict[RenderDictkey][0] == True:
                    # 将该单元格所在行的第二列的文字变为√
                    table.cell(RowNumberDict[Tablekey],1).text = "√"
                # 如果字典中的key对应的value的第一个元素为False
                if RenderDict[RenderDictkey][0] == False:
                    # 将该单元格所在行的第三列的文字变为√
                    table.cell(RowNumberDict[Tablekey],2).text = "√"
                # 将该单元格所在行第四列的文字变为字典中的value的第二个元素
                table.cell(RowNumberDict[Tablekey],3).text = RenderDict[RenderDictkey][1]
                
    TemplateFileName += "（" + "已填充完成" + "）.docx"
    # 保存文件
    Doc.save(OutputDir + "\\" + TemplateFileName)
    return